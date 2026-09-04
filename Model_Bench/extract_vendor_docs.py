#!/usr/bin/env python3
"""Extract the real content of the Sohar vendor .docx handover documents
into plain markdown -- paragraphs AND tables (tables carry the real
technical detail: tag-name mappings, per-state configuration, condition
logic -- paragraphs alone are just section headers and overview prose).

Why this exists: these 5 documents (EAF/LRF/CCM Per Heat Event, SMS Plant
Process Time, Billets Cast Count -- real vendor handover docs, "SITC of
X-Force Historian and Dashboard Development", Sohar Steel Oman) were
sitting in Reference Documents/Sohar_Vendor_Docs/ unread this entire
project until 2026-09-03, despite being real, given source material for
exactly the domain (sohar-sms-event-workflows.md) this project already
tries to document from live SQL investigation alone. This closes that
gap by pulling the real content out in a form Knowledge/ docs and skills
can actually cite.

Usage:
    python extract_vendor_docs.py
"""
import sys
from pathlib import Path

import docx

VENDOR_DIR = Path(r"C:\Users\Admin\Documents\Office\AIHelpdesk\Reference Documents\Sohar_Vendor_Docs")
OUTPUT_DIR = Path(r"C:\Users\Admin\Documents\Office\AIHelpdesk\Knowledge\vendor_docs_extracted")


def extract_docx(path: Path) -> str:
    d = docx.Document(str(path))
    lines = [f"# {path.stem}", "", f"*Extracted from `{path.name}` -- real vendor handover doc, not invented.*", ""]

    # Interleave paragraphs and tables in document order where possible.
    # python-docx doesn't give a clean unified iterator, so approximate:
    # emit all paragraphs first (headings/prose), then all tables in order
    # (numbered so they can be cross-referenced back to the source doc).
    for p in d.paragraphs:
        if p.text.strip():
            style = (p.style.name or "").lower()
            if "heading" in style or "title" in style:
                level = "#" * min(int("".join(c for c in style if c.isdigit()) or "2"), 4)
                lines.append(f"\n{level} {p.text.strip()}\n")
            else:
                lines.append(p.text.strip())

    lines.append("\n---\n## Tables (real technical detail -- tag mappings, state configs, condition logic)\n")
    for i, t in enumerate(d.tables):
        lines.append(f"\n### Table {i + 1}\n")
        rows = [[c.text.strip() for c in row.cells] for row in t.rows]
        if not rows:
            continue
        header = rows[0]
        lines.append("| " + " | ".join(header) + " |")
        lines.append("|" + "---|" * len(header))
        for row in rows[1:]:
            row = row + [""] * (len(header) - len(row))  # pad short rows
            lines.append("| " + " | ".join(c.replace("\n", " ") for c in row[:len(header)]) + " |")

    return "\n".join(lines)


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    docx_files = sorted(VENDOR_DIR.glob("*.docx"))
    if not docx_files:
        print(f"No .docx files found in {VENDOR_DIR}")
        sys.exit(1)

    for path in docx_files:
        print(f"Extracting {path.name}...")
        content = extract_docx(path)
        out_path = OUTPUT_DIR / f"{path.stem}.md"
        out_path.write_text(content, encoding="utf-8")
        print(f"  -> {out_path} ({len(content)} chars)")

    print(f"\nDone. {len(docx_files)} document(s) extracted to {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
